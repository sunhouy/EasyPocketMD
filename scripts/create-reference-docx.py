#!/usr/bin/env python3
"""
创建Pandoc reference.docx模板，设置中文字体
"""
import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("错误: 需要安装python-docx库")
    print("运行: pip install python-docx")
    sys.exit(1)


def create_reference_docx(output_path, title_font='SimHei', body_font='SimSun'):
    """创建reference.docx模板"""

    # 创建新文档
    doc = Document()

    # 修改样式
    styles = doc.styles

    # 标题样式 (Heading 1-6)
    for i in range(1, 7):
        style_name = f'Heading {i}'
        try:
            style = styles[style_name]
        except KeyError:
            # 如果样式不存在，创建它
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        font = style.font
        font.name = title_font
        font.color.rgb = RGBColor(0, 0, 0)
        font.bold = True
        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), title_font)

        # 设置字号
        sizes = {1: 26, 2: 22, 3: 18, 4: 16, 5: 14, 6: 12}
        font.size = Pt(sizes.get(i, 12))

    # 正文样式
    try:
        style = styles['Normal']
    except KeyError:
        style = styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)

    font = style.font
    font.name = body_font
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), body_font)

    # 列表样式
    for style_name in ['List Paragraph', 'List', 'List Bullet', 'List Number']:
        try:
            style = styles[style_name]
            font = style.font
            font.name = body_font
            font.color.rgb = RGBColor(0, 0, 0)
            style.element.rPr.rFonts.set(qn('w:eastAsia'), body_font)
        except KeyError:
            pass  # 样式不存在，跳过

    # 添加示例内容
    doc.add_heading('一级标题', level=1)
    doc.add_heading('二级标题', level=2)
    doc.add_heading('三级标题', level=3)
    doc.add_paragraph('这是正文段落。这是正文段落。这是正文段落。')
    doc.add_paragraph('列表项1', style='List Bullet')
    doc.add_paragraph('列表项2', style='List Bullet')

    # 保存文档
    doc.save(str(output_path))
    return True


def main():
    script_dir = Path(__file__).parent
    project_dir = script_dir.parent
    reference_dir = project_dir / 'reference-docs'
    reference_dir.mkdir(exist_ok=True)

    # 创建默认模板
    reference_path = reference_dir / 'reference.docx'
    print(f"创建reference.docx模板: {reference_path}")

    if create_reference_docx(reference_path, 'SimHei', 'SimSun'):
        print(f"✓ 成功创建: {reference_path}")
        print(f"\n在.env文件中添加:")
        print(f"PANDOC_REFERENCE_DOCX={reference_path}")

        # 创建不同字体组合的模板
        variants = [
            ('reference-simhei-simsun.docx', 'SimHei', 'SimSun'),
            ('reference-simhei-simkai.docx', 'SimHei', 'SimKai'),
            ('reference-simsun-simsun.docx', 'SimSun', 'SimSun'),
            ('reference-simkai-simsun.docx', 'SimKai', 'SimSun'),
        ]

        print("\n创建其他字体组合模板...")
        for filename, title_font, body_font in variants:
            path = reference_dir / filename
            if create_reference_docx(path, title_font, body_font):
                print(f"✓ {filename} (标题:{title_font}, 正文:{body_font})")
    else:
        print("✗ 创建失败")
        sys.exit(1)


if __name__ == '__main__':
    main()
