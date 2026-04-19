#!/usr/bin/env python3
"""
根据用户选择的字体和字号动态生成DOCX模板
"""
import os
import sys
from pathlib import Path
import tempfile
import json

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("错误: 需要安装python-docx库")
    print("运行: pip install python-docx")
    sys.exit(1)


def generate_docx_template(output_path, title_font='SimHei', body_font='SimSun', 
                         title_font_size=24, body_font_size=12, 
                         h1_size=32, h2_size=28, h3_size=24, 
                         h4_size=20, h5_size=18, h6_size=16):
    """根据用户设置生成DOCX模板"""

    # 创建新文档
    doc = Document()

    # 修改样式
    styles = doc.styles

    # 标题样式 (Heading 1-6)
    heading_sizes = {
        1: h1_size,
        2: h2_size,
        3: h3_size,
        4: h4_size,
        5: h5_size,
        6: h6_size
    }

    for i in range(1, 7):
        style_name = f'Heading {i}'
        try:
            style = styles[style_name]
        except KeyError:
            # 如果样式不存在，创建它
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        font = style.font
        font.name = title_font
        font.color.rgb = RGBColor(0, 0, 0)
        font.bold = True
        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), title_font)

        # 设置字号
        font.size = Pt(heading_sizes.get(i, 12))

    # 正文样式
    try:
        style = styles['Normal']
    except KeyError:
        style = styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)

    font = style.font
    font.name = body_font
    font.size = Pt(body_font_size)
    font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), body_font)

    # 列表样式
    for style_name in ['List Paragraph', 'List', 'List Bullet', 'List Number']:
        try:
            style = styles[style_name]
            font = style.font
            font.name = body_font
            font.size = Pt(body_font_size)
            font.color.rgb = RGBColor(0, 0, 0)
            style.element.rPr.rFonts.set(qn('w:eastAsia'), body_font)

        except KeyError:
            pass  # 样式不存在，跳过

    # 添加示例内容
    doc.add_heading('一级标题', level=1)
    doc.add_heading('二级标题', level=2)
    doc.add_heading('三级标题', level=3)
    doc.add_paragraph('这是正文段落。这是正文段落。这是正文段落。')
    doc.add_paragraph('列表项1', style='List Bullet')
    doc.add_paragraph('列表项2', style='List Bullet')

    # 保存文档
    doc.save(str(output_path))
    return True


def main():
    """主函数，处理命令行参数"""
    if len(sys.argv) != 2:
        print("用法: python generate-docx-template.py '<config_json>'")
        sys.exit(1)

    try:
        config = json.loads(sys.argv[1])
    except json.JSONDecodeError as e:
        print(f"错误: 无效的JSON配置: {e}")
        sys.exit(1)

    # 提取配置参数
    title_font = config.get('titleFont', 'SimHei')
    body_font = config.get('bodyFont', 'SimSun')
    title_font_size = float(config.get('titleFontSize', 24))
    body_font_size = float(config.get('bodyFontSize', 12))
    h1_size = float(config.get('h1Size', title_font_size * 2))
    h2_size = float(config.get('h2Size', title_font_size * 1.6))
    h3_size = float(config.get('h3Size', title_font_size * 1.35))
    h4_size = float(config.get('h4Size', title_font_size * 1.2))
    h5_size = float(config.get('h5Size', title_font_size))
    h6_size = float(config.get('h6Size', max(14, title_font_size * 0.9)))

    # 生成临时文件路径
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        temp_path = tmp.name

    try:
        # 生成模板
        if generate_docx_template(
            temp_path, 
            title_font=title_font,
            body_font=body_font,
            title_font_size=title_font_size,
            body_font_size=body_font_size,
            h1_size=h1_size,
            h2_size=h2_size,
            h3_size=h3_size,
            h4_size=h4_size,
            h5_size=h5_size,
            h6_size=h6_size
        ):
            # 输出临时文件路径
            print(temp_path)
        else:
            print("错误: 生成模板失败")
            sys.exit(1)
    except Exception as e:
        print(f"错误: {e}")
        if os.path.exists(temp_path):
            os.unlink(temp_path)
        sys.exit(1)


if __name__ == '__main__':
    main()
